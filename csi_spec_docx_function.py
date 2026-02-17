# Add this import at the top of drex_calculator.py:
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Replace the generate_csi_specification() function with this:

def generate_csi_specification_docx(project_info, dryers, manifold_info, results):
    """Generate CSI specification as Word document"""
    doc = Document()
    
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("SECTION: 23 35 01")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("COMMERCIAL DRYER EXHAUST SYSTEM")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # PART 1: GENERAL
    heading1 = doc.add_heading("PART 1: GENERAL", level=1)
    
    # 1.01 SUMMARY
    doc.add_heading("1.01 SUMMARY", level=2)
    p = doc.add_paragraph()
    p.add_run("A. This system shall provide the required exhaust capacity for commercial clothes dryer applications. The following are components of the system:")
    
    dryers_list = doc.add_paragraph(style='List Number')
    dryers_list.add_run(f"LF Systems, Commercial Dryer Exhaust Fan Model DEF{results.get('selected_fan', {}).get('model', 'XX').replace('DEF', '')}, ETL-listed to UL STD 705.")
    
    controls_list = doc.add_paragraph(style='List Number')
    controls_list.add_run("LF Systems Model L150 Constant Pressure Controller")
    
    elec_list = doc.add_paragraph(style='List Number')
    elec_list.add_run("Electrical connections, by installing contractor.")
    
    # 1.02 RELATED REQUIREMENTS  
    doc.add_heading("1.02 RELATED REQUIREMENTS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Section 01 41 00, REGULATORY REQUIREMENTS")
    
    # 1.03 CODES AND STANDARDS
    doc.add_heading("1.03 CODES AND STANDARDS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. The following published specification standards apply to work in this section:")
    
    for code in ["UL -- Underwriters Laboratories", "ANSI Z223.1", "NFPA 211", "International Mechanical Code", "National Electrical Code"]:
        code_p = doc.add_paragraph(style='List Number')
        code_p.add_run(code)
    
    # 1.04 SUBMITTALS
    doc.add_heading("1.04 SUBMITTALS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Submittal documents shall be provided by the local manufacturer's representative.")
    
    p = doc.add_paragraph()
    p.add_run("B. The following shall be submitted to the Owner's representative:")
    
    submit_items = [
        "All relevant data sheets for the dryer exhaust fan, controller and related controls.",
        "ASHRAE duct design calculations.",
        "Wiring diagrams and installation manuals shall be submitted to the installing contractor prior to installation.",
        "Certification of listing for the actual application by an OSHA approved NRTL."
    ]
    
    for item in submit_items:
        item_p = doc.add_paragraph(style='List Number')
        item_p.add_run(item)
    
    # 1.05 QUALITY ASSURANCE
    doc.add_heading("1.05 QUALITY ASSURANCE", level=2)
    p = doc.add_paragraph()
    p.add_run(f"A. All listed dryer exhaust fans shall be assembled and marked at the facility indicated by the listing control number.")
    
    p = doc.add_paragraph()
    p.add_run("B. Exhaust system must be able to sense pressure in the venting system and maintain constant pressure.")
    
    p = doc.add_paragraph()
    p.add_run(f"C. System capacity scheduled shall be minimum {results.get('total_cfm', 0):.0f} CFM at {results.get('total_system_dp', 0):.2f} inches water column static pressure.")
    
    # 1.06 WARRANTY
    doc.add_heading("1.06 MANUFACTURER WARRANTY", level=2)
    p = doc.add_paragraph()
    p.add_run("A. All equipment is to be guaranteed against defects in materials and/or workmanship for a period of 24 months from the date of delivery to the construction site.")
    
    # PART 2: PRODUCTS
    doc.add_page_break()
    heading2 = doc.add_heading("PART 2: PRODUCTS", level=1)
    
    # 2.01 MANUFACTURER
    doc.add_heading("2.01 MANUFACTURER, DRYER EXHAUST SYSTEM", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Furnish LF Systems Commercial Dryer Exhaust System consisting of the following components:")
    
    fan_model = results.get('selected_fan', {}).get('model', 'DEFXX')
    components = [
        f"Model {fan_model} Inline dryer exhaust fan with design volume of {results.get('total_cfm', 0):.0f} CFM and design pressure of {results.get('total_system_dp', 0):.2f} inches water column as scheduled.",
        "Model L150 constant pressure controller with integrated display and alarm outputs.",
        "Current sensing relay (if required for application)."
    ]
    
    for comp in components:
        comp_p = doc.add_paragraph(style='List Number')
        comp_p.add_run(comp)
    
    # 2.02 EXHAUST FAN DESCRIPTION
    doc.add_heading("2.02 DESCRIPTION, INLINE DRYER EXHAUST FAN", level=2)
    
    fan_desc = [
        "The entire dryer exhaust fan shall be constructed of G90 galvanized steel with minimum .030 thickness. The fan shall be listed to UL STD 705 and shall bear the listed mark from an OSHA approved NRTL.",
        "The fan impeller shall be statically and dynamically balanced with permanently attached balancing weights of the same material as the impeller.",
        "The dryer exhaust fan shall be listed for 480°F exhaust gas temperatures and shall discharge gases as scheduled on drawings. The fan shall include integrated access panel.",
        "The fan motor shall be electronically commutated (EC), totally enclosed and outdoor rated with minimum efficiency of 75%. The motor shall provide variable speed control for pressure management."
    ]
    
    for idx, desc in enumerate(fan_desc, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # 2.03 CONTROLLER DESCRIPTION
    doc.add_heading("2.03 DESCRIPTION, L150 CONSTANT PRESSURE CONTROLLER", level=2)
    
    controller_desc = [
        "The L150 constant pressure controller shall monitor and control pressure in the dryer exhaust duct to maintain programmed setpoint pressure.",
        "The controller's pressure range shall be -1.00\" w.c. to +1.00\" w.c. with resolution of 0.01\" w.c.",
        "The controller shall provide visual LCD display, audible alarm, and dry contact alarm outputs for remote monitoring.",
        f"The controller shall maintain setpoint within +/-0.01\" w.c. Setpoint for this application: {results.get('manifold_dp', 0):.2f}\" w.c.",
        "The L150 shall include sensor verification, sleep mode, and adjustable PID control parameters."
    ]
    
    for idx, desc in enumerate(controller_desc, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # 2.04 PERFORMANCE
    doc.add_heading("2.04 PERFORMANCE REQUIREMENTS", level=2)
    
    perf = [
        "The dryer exhaust fan system shall reach setpoint within 15 seconds of initial demand.",
        f"Fan shall deliver minimum {results.get('total_cfm', 0):.0f} CFM at {results.get('total_system_dp', 0):.2f}\" w.c. system static pressure.",
        "The system shall include intelligent feedback signal to determine and control motor RPM.",
        "System shall operate quietly with vibration isolation as required."
    ]
    
    for idx, desc in enumerate(perf, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # 2.05 SEQUENCE OF OPERATION
    doc.add_heading("2.05 SEQUENCE OF OPERATION", level=2)
    
    seq = [
        "Upon demand signal, the L150 controller activates sensor verification. Once verified, the system controls fan speed to achieve setpoint pressure.",
        "As individual dryers call for heat, the L150 adjusts fan speed to maintain constant setpoint pressure in the exhaust manifold.",
        "When all dryers have satisfied, the controller engages sleep mode after programmable delay period.",
        "System includes integrated alarm functions for low pressure, high pressure, and sensor fault conditions."
    ]
    
    for idx, desc in enumerate(seq, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # 2.06 ELECTRICAL
    doc.add_heading("2.06 ELECTRICAL REQUIREMENTS", level=2)
    
    p = doc.add_paragraph()
    p.add_run("A. Power supply shall be:")
    
    elec = [
        "L150 pressure controller: 120VAC, single phase, 60Hz",
        "EC motor fan: 120VAC or 240VAC single phase, 60Hz (as scheduled)",
        "All wiring in accordance with National Electrical Code"
    ]
    
    for item in elec:
        elec_p = doc.add_paragraph(style='List Number')
        elec_p.add_run(item)
    
    # PART 3: EXECUTION
    doc.add_page_break()
    heading3 = doc.add_heading("PART 3: EXECUTION", level=1)
    
    # 3.01 INSTALLATION
    doc.add_heading("3.01 INSTALLATION", level=2)
    
    install = [
        "Complete structural, mechanical, and electrical connections in accordance with manufacturer's printed instructions.",
        f"Install {len(dryers)} dryer connectors to manifold duct of {manifold_info.get('diameter', 'XX')}\" diameter as shown on drawings.",
        f"Install manifold duct system of {manifold_info.get('length', 'XX')} feet total length with fittings as scheduled.",
        f"Mount {fan_model} exhaust fan in location shown on drawings with proper vibration isolation and support.",
        "Install L150 controller in accessible location for programming and service. Mount pressure sensor in manifold as detailed.",
        "Low voltage control wiring shall be minimum 18 AWG shielded wire. Maintain separation from line voltage wiring.",
        "Provide electrical disconnects and overcurrent protection as required by NEC."
    ]
    
    for idx, desc in enumerate(install, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # 3.02 STARTUP
    doc.add_heading("3.02 STARTUP AND COMMISSIONING", level=2)
    
    startup = [
        "System startup shall be performed by qualified LF Systems technician or authorized representative.",
        "Verify all electrical connections, duct connections, and sensor locations prior to energizing system.",
        f"Program L150 controller with setpoint pressure of {results.get('manifold_dp', 0):.2f}\" w.c. and verify sensor calibration.",
        "Test system operation with all connected dryers and verify pressure control performance.",
        "Provide training to owner's maintenance personnel on system operation, troubleshooting, and routine maintenance."
    ]
    
    for idx, desc in enumerate(startup, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    
    # Footer
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("LF SYSTEMS by RM Manifold Group Inc.\n")
    footer_run.bold = True
    footer_run.font.size = Pt(12)
    footer_run2 = footer_p.add_run("www.lfsystems.net\n")
    footer_run2.font.size = Pt(10)
    footer_run3 = footer_p.add_run(f"Project: {project_info.get('name', 'Commercial Dryer Exhaust')}\n")
    footer_run3.font.size = Pt(10)
    footer_run4 = footer_p.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    footer_run4.font.size = Pt(10)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# END OF FUNCTION
