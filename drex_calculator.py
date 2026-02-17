"""
DREX - Dryer Exhaust Calculator
Commercial Dryer Exhaust System Sizing Program
Developed for LF Systems by RM Manifold
"""

import streamlit as st
import pandas as pd
import numpy as np
import requests
from datetime import datetime
import json
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
import base64
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Constants
DRYER_CONNECTOR_MAX_DP = 0.25  # IN WC
MANIFOLD_MAX_DP = 1.0  # IN WC
AIR_DENSITY = 0.0696  # lb/ft^3 at 120°F (dryer exhaust temperature)
DUCT_FRICTION_FACTOR = 0.35  # Friction factor for duct

# K-values for fittings (dimensionless) - Updated per engineering standards
K_VALUES = {
    "90° Elbow": 0.5,
    "45° Elbow": 0.25,
    "30° Elbow": 0.15,
    "Lateral Tee": 0.75,
    "Entry (flush)": 0.5,
    "Exit": 1.0
}

# DEF Fan Curve Data
DEF_FAN_CURVES = {
    "DEF04": {
        "CFM": [540, 490, 430, 350, 240],
        "SP": [0, 0.25, 0.5, 0.75, 1.0]
    },
    "DEF08": {
        "CFM": [970, 890, 840, 780, 680, 540, 440, 270],
        "SP": [0, 0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75]
    },
    "DEF015": {
        "CFM": [1860, 1780, 1700, 1610, 1520, 1410, 1280, 1140, 990],
        "SP": [0, 0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0]
    },
    "DEF025": {
        "CFM": [2480, 2400, 2320, 2230, 2140, 2040, 1930, 1790, 1630],
        "SP": [0, 0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0]
    },
    "DEF035": {
        "CFM": [4100, 3940, 3770, 3610, 3460, 3300, 3120, 2900, 2630],
        "SP": [0, 0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0]
    },
    "DEF050": {
        "CFM": [5850, 5660, 5450, 5300, 5090, 4890, 4680, 4460, 4230],
        "SP": [0, 0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0]
    }
}

# LF Systems Branding Colors (from RM Manifold Brand Guide)
BRAND_PRIMARY = "#2a3853"  # Pantone 533C
BRAND_SECONDARY = "#234699"  # Pantone 7686C  
BRAND_ACCENT = "#b11f33"  # Pantone 187C
BRAND_GRAY = "#97999b"  # Cool Gray 7C

def plot_fan_and_system_curves(selected_fan_model, required_cfm, total_system_sp, manifold_sp):
    """
    Generate fan curve and system curve plot
    
    Args:
        selected_fan_model: DEF fan model name
        required_cfm: Total system CFM
        total_system_sp: Total system pressure (for operating point marker)
        manifold_sp: Manifold pressure only (for system curve calculation)
    """
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Plot fan curve
    fan_data = DEF_FAN_CURVES[selected_fan_model]
    ax.plot(fan_data['CFM'], fan_data['SP'], 'b-', linewidth=2.5, label=f'{selected_fan_model} Fan Curve', color=BRAND_SECONDARY)
    ax.scatter(fan_data['CFM'], fan_data['SP'], color=BRAND_SECONDARY, s=60, zorder=5, edgecolors='white', linewidths=1.5)
    
    # Plot system curve using MANIFOLD PRESSURE ONLY (parabolic: SP = k * CFM^2)
    k = manifold_sp / (required_cfm ** 2) if required_cfm > 0 else 0
    cfm_range = np.linspace(0, max(fan_data['CFM']) * 1.1, 100)
    system_curve = k * (cfm_range ** 2)
    ax.plot(cfm_range, system_curve, '--', linewidth=2.5, label='System Curve (Manifold Only)', color=BRAND_ACCENT)
    
    # Mark operating point on the system curve (using manifold pressure only)
    # This shows where the system operates on the manifold resistance curve
    ax.scatter([required_cfm], [manifold_sp], color='#00a86b', s=250, marker='*', 
               zorder=10, label=f'Operating Point\n({required_cfm:.0f} CFM @ {manifold_sp:.2f}" WC Manifold)\nTotal System: {total_system_sp:.2f}" WC',
               edgecolors='white', linewidths=2)
    
    # Formatting
    ax.set_xlabel('Airflow (CFM)', fontsize=13, fontweight='bold')
    ax.set_ylabel('Static Pressure (IN WC)', fontsize=13, fontweight='bold')
    ax.set_title(f'Fan Performance & System Curve - {selected_fan_model}', fontsize=15, fontweight='bold', pad=20)
    ax.grid(True, alpha=0.3, linestyle=':', linewidth=1)
    ax.legend(loc='upper right', fontsize=11, framealpha=0.95)
    ax.set_xlim(0, max(fan_data['CFM']) * 1.15)
    ax.set_ylim(0, max(fan_data['SP']) * 1.15)
    
    # Add LF Systems watermark
    ax.text(0.98, 0.02, 'LF Systems by RM Manifold | www.lfsystems.net', 
            transform=ax.transAxes, fontsize=9, color='gray', alpha=0.7,
            ha='right', va='bottom', style='italic')
    
    plt.tight_layout()
    return fig

def initialize_session_state():
    """Initialize session state variables"""
    if 'project_info' not in st.session_state:
        st.session_state.project_info = {}
    if 'dryers' not in st.session_state:
        st.session_state.dryers = []
    if 'manifold_info' not in st.session_state:
        st.session_state.manifold_info = {}
    if 'step' not in st.session_state:
        st.session_state.step = 'welcome'
    if 'calculation_results' not in st.session_state:
        st.session_state.calculation_results = {}

def get_elevation_from_zip(zip_code):
    """Get elevation from zip code using external API"""
    try:
        # Get location from zip code
        response = requests.get(f"https://api.zippopotam.us/us/{zip_code}", timeout=5)
        if response.status_code == 200:
            data = response.json()
            lat = float(data['places'][0]['latitude'])
            lon = float(data['places'][0]['longitude'])
            city = data['places'][0]['place name']
            state = data['places'][0]['state abbreviation']
            
            # Get elevation
            elev_response = requests.get(
                f"https://api.open-elevation.com/api/v1/lookup?locations={lat},{lon}",
                timeout=5
            )
            if elev_response.status_code == 200:
                elevation = elev_response.json()['results'][0]['elevation']
                elevation_ft = elevation * 3.28084  # Convert meters to feet
                return {
                    'city': city,
                    'state': state,
                    'elevation': round(elevation_ft, 0),
                    'success': True
                }
    except:
        pass
    return {'success': False}

def calculate_duct_pressure_loss(length, diameter, velocity, k_sum, rho=AIR_DENSITY):
    """
    Calculate pressure loss using the duct design equation
    dp = ((0.35*L/D) + SUM(k)) * rho * (V/1096.2)^2
    
    Where:
        0.35 = friction factor for duct
        L = length in feet
        D = diameter in inches
        k = sum of fitting loss coefficients (dimensionless)
        rho = air density at 120°F (0.0696 lb/ft^3)
        V = velocity in feet per minute
    
    Args:
        length: duct length in feet
        diameter: duct diameter in inches
        velocity: velocity in feet per minute
        k_sum: sum of all k-values (dimensionless)
        rho: air density in lb/ft^3 (default: 0.0696 at 120°F)
    
    Returns:
        pressure loss in IN WC
    """
    if diameter <= 0:
        return 0
    
    # Note: diameter is in inches, so L/D gives friction in feet/inch
    friction_term = DUCT_FRICTION_FACTOR * (length / diameter)
    total_term = friction_term + k_sum
    velocity_term = (velocity / 1096.2) ** 2
    
    dp = total_term * rho * velocity_term
    return dp

def calculate_velocity(cfm, diameter):
    """Calculate velocity in FPM given CFM and diameter in inches"""
    area_sq_ft = (np.pi * (diameter / 12) ** 2) / 4
    if area_sq_ft > 0:
        return cfm / area_sq_ft
    return 0

def optimize_manifold_diameter(total_cfm, length, k_sum, max_dp=MANIFOLD_MAX_DP):
    """
    Optimize manifold diameter to achieve maximum pressure loss of max_dp
    """
    # Standard duct diameters in inches
    standard_diameters = [4, 5, 6, 7, 8, 9, 10, 12, 14, 16, 18, 20, 22, 24, 26, 28, 30, 32, 34, 36]
    
    for diameter in standard_diameters:
        velocity = calculate_velocity(total_cfm, diameter)
        dp = calculate_duct_pressure_loss(length, diameter, velocity, k_sum)
        if dp <= max_dp:
            return diameter, velocity, dp
    
    # If no standard diameter works, return the largest
    diameter = standard_diameters[-1]
    velocity = calculate_velocity(total_cfm, diameter)
    dp = calculate_duct_pressure_loss(length, diameter, velocity, k_sum)
    return diameter, velocity, dp

def select_def_fan(required_cfm, required_sp):
    """Select appropriate DEF fan based on CFM and static pressure requirements"""
    suitable_fans = []
    
    for fan_model, curve_data in DEF_FAN_CURVES.items():
        cfm_values = curve_data['CFM']
        sp_values = curve_data['SP']
        
        # Interpolate to find available CFM at required SP
        if required_sp <= max(sp_values):
            available_cfm = np.interp(required_sp, sp_values, cfm_values)
            if available_cfm >= required_cfm:
                # Calculate margin
                margin = ((available_cfm - required_cfm) / required_cfm) * 100
                suitable_fans.append({
                    'model': fan_model,
                    'available_cfm': available_cfm,
                    'margin': margin,
                    'max_cfm': max(cfm_values),
                    'max_sp': max(sp_values)
                })
    
    # Sort by margin (prefer fans with 10-25% margin)
    suitable_fans.sort(key=lambda x: abs(x['margin'] - 17.5))
    
    return suitable_fans

def generate_pdf_report(project_info, dryers, manifold_info, results):
    """Generate professional sales-focused PDF report with LF Systems branding"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        topMargin=0.75*inch, 
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    story = []
    styles = getSampleStyleSheet()
    # Custom Styles with LF Systems branding
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2a3853'),  # LF Systems primary color
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor('#234699'),  # LF Systems secondary
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2a3853'),
        spaceBefore=16,
        spaceAfter=10,
        fontName='Helvetica-Bold',
        borderPadding=5,
        leftIndent=0
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.black,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica'
    )
    # LF Systems Logo (if available)
    try:
        logo = Image('lf_systems_logo.jpg', width=3*inch, height=0.75*inch)
        logo.hAlign = 'CENTER'
        story.append(logo)
        story.append(Spacer(1, 0.2*inch))
    except:
        # Fallback if logo not available
        logo_text = Paragraph("<b>LF SYSTEMS</b><br/>by RM Manifold Group Inc.", subtitle_style)
        story.append(logo_text)
        story.append(Spacer(1, 0.2*inch))
    # Title
    title = Paragraph("Commercial Dryer Exhaust<br/>System Proposal", title_style)
    story.append(title)
    subtitle = Paragraph(f"<i>{project_info.get('name', 'Your Project')}</i>", subtitle_style)
    story.append(subtitle)
    # Project Info Box
    project_data = [
        ['Project Location:', f"{project_info.get('city', 'N/A')}, {project_info.get('state', 'N/A')}"],
        ['Elevation:', f"{project_info.get('elevation', 'N/A')} ft"],
        ['Prepared For:', project_info.get('user_name', 'N/A')],
        ['Date:', datetime.now().strftime('%B %d, %Y')],
    ]
    project_table = Table(project_data, colWidths=[2*inch, 4*inch])
    project_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f2f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    story.append(project_table)
    story.append(Spacer(1, 0.3*inch))
    # Executive Summary Section
    story.append(Paragraph("Executive Summary", heading_style))
    exec_summary = f"""
    LF Systems is pleased to provide this comprehensive dryer exhaust system solution for your facility. 
    Our engineered approach ensures optimal performance, code compliance, and long-term reliability.
    <br/><br/>
    <b>System Highlights:</b><br/>
    • {len(dryers)} commercial dryers supported<br/>
    • {results['total_cfm']:.0f} CFM total exhaust capacity<br/>
    • {manifold_info['diameter']}" manifold duct system<br/>
    • DEF Series exhaust fan with L150 intelligent controls<br/>
    • Fully code-compliant design meeting IMC and NFPA standards
    """
    story.append(Paragraph(exec_summary, body_style))
    story.append(Spacer(1, 0.2*inch))
    # Key System Specifications - Highlighted Box
    story.append(Paragraph("Recommended System Configuration", heading_style))
    spec_data = [
        ['SPECIFICATION', 'VALUE'],
        ['Total System Airflow', f"{results['total_cfm']:.0f} CFM"],
        ['Manifold Diameter', f"{manifold_info['diameter']}\""],
        ['Manifold Length', f"{manifold_info['length']} feet"],
        ['System Static Pressure', f"{results['total_system_dp']:.2f}\" WC"],
        ['Recommended Fan', results['selected_fan']['model'] if results['selected_fan'] else 'Contact Us'],
        ['Controller', 'L150 Constant Pressure'],
    ]
    spec_table = Table(spec_data, colWidths=[3*inch, 2.5*inch])
    spec_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#234699')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#234699')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
    ]))
    story.append(spec_table)
    story.append(Spacer(1, 0.3*inch))
    # Benefits Section
    story.append(Paragraph("Why Choose LF Systems?", heading_style))
    benefits = """
    <b>Industry-Leading Performance:</b> Our DEF Series fans deliver consistent, reliable exhaust 
    with intelligent pressure control that adapts to changing load conditions.<br/><br/>
    <b>Code Compliance Guaranteed:</b> Every system is engineered to meet or exceed IMC, NFPA 211, 
    and local building code requirements. Our UL-listed components ensure safety and reliability.<br/><br/>
    <b>Energy Efficiency:</b> Variable-speed EC motors and intelligent controls minimize energy 
    consumption while maintaining optimal exhaust performance.<br/><br/>
    <b>Proven Reliability:</b> Backed by RM Manifold Group's decades of HVAC expertise and a 
    comprehensive 24-month warranty on all equipment.
    """
    story.append(Paragraph(benefits, body_style))
    story.append(PageBreak())
    # Technical Details Section
    story.append(Paragraph("System Design Details", heading_style))
    # Dryer Configuration
    story.append(Paragraph("<b>Dryer Configuration:</b>", body_style))
    dryer_data = [['#', 'CFM', 'Outlet Ø', 'Connector', 'Pressure Loss']]
    for idx, dryer in enumerate(dryers[:10], 1):  # Show first 10
        dryer_data.append([
            str(idx),
            f"{dryer['cfm']} CFM",
            f"{dryer['outlet_diameter']}\"",
            f"{dryer['connector_length']} ft",
            f"{dryer['connector_dp']:.3f}\" WC"
        ])
    if len(dryers) > 10:
        dryer_data.append(['...', f'+{len(dryers)-10} more', '', '', ''])
    dryer_table = Table(dryer_data, colWidths=[0.5*inch, 1.2*inch, 1*inch, 1.2*inch, 1.3*inch])
    dryer_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2a3853')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
    ]))
    story.append(dryer_table)
    story.append(Spacer(1, 0.2*inch))
    # Manifold System
    story.append(Paragraph("<b>Manifold System:</b>", body_style))
    manifold_text = f"""
    • Length: {manifold_info['length']} feet<br/>
    • Diameter: {manifold_info['diameter']} inches<br/>
    • Fittings: {manifold_info.get('fittings_summary', 'As specified')}<br/>
    • Velocity: {results['manifold_velocity']:.0f} FPM<br/>
    • Pressure Loss: {results['manifold_dp']:.3f}\" WC (within {MANIFOLD_MAX_DP}\" WC limit)
    """
    story.append(Paragraph(manifold_text, body_style))
    story.append(Spacer(1, 0.2*inch))
    # Fan Selection
    if results['selected_fan']:
        fan = results['selected_fan']
        story.append(Paragraph("<b>Exhaust Fan Selection:</b>", body_style))
        fan_info = f"""
        <b>Model:</b> {fan['model']}<br/>
        <b>Available CFM at Operating Point:</b> {fan['available_cfm']:.0f} CFM<br/>
        <b>Design Margin:</b> {fan['margin']:.1f}% (ensures reliable performance)<br/>
        <b>Maximum Capacity:</b> {fan['max_cfm']:.0f} CFM @ {fan['max_sp']:.2f}\" WC
        """
        story.append(Paragraph(fan_info, body_style))
        # Status indicator
        if fan['margin'] >= 10:
            status_text = '<font color="#00a86b">✓ EXCELLENT - Fan properly sized with adequate safety margin</font>'
        else:
            status_text = '<font color="#b11f33">⚠ Contact LF Systems for optimization recommendations</font>'
        story.append(Paragraph(status_text, body_style))
    story.append(Spacer(1, 0.3*inch))
    # Engineering Notes
    story.append(Paragraph("Engineering Notes", heading_style))
    eng_notes = f"""
    <b>Design Standards:</b><br/>
    • Calculations per ASHRAE Fundamentals<br/>
    • Pressure loss: ΔP = (0.35 × L/D + ΣK) × ρ × (V/1096.2)²<br/>
    • Air density: 0.0696 lb/ft³ @ 120°F (dryer exhaust temperature)<br/>
    • Code compliance: IMC, NFPA 211, UL 705<br/><br/>
    <b>System Warnings:</b><br/>
    {"• ⚠ Connector pressure loss exceeds 0.25\" WC recommendation<br/>" if results['worst_connector_dp'] > DRYER_CONNECTOR_MAX_DP else ""}
    {"• ⚠ Manifold pressure loss exceeds 1.0\" WC maximum<br/>" if results['manifold_dp'] > MANIFOLD_MAX_DP else ""}
    {"• ✓ All pressure losses within acceptable limits<br/>" if results['manifold_dp'] <= MANIFOLD_MAX_DP and results['worst_connector_dp'] <= DRYER_CONNECTOR_MAX_DP else ""}
    """
    story.append(Paragraph(eng_notes, body_style))
    # New page for Next Steps
    story.append(PageBreak())
    # Next Steps / Call to Action
    story.append(Paragraph("Next Steps", heading_style))
    next_steps = """
    <b>1. Review & Approval</b><br/>
    Review this proposal and confirm it meets your project requirements.<br/><br/>
    <b>2. Detailed Engineering</b><br/>
    Our engineering team will provide detailed submittal drawings and specifications.<br/><br/>
    <b>3. Equipment Procurement</b><br/>
    Once approved, equipment will be manufactured and shipped to your project site.<br/><br/>
    <b>4. Installation Support</b><br/>
    LF Systems provides complete installation support and startup assistance.<br/><br/>
    <b>5. Training & Warranty</b><br/>
    Comprehensive operator training and 24-month equipment warranty included.
    """
    story.append(Paragraph(next_steps, body_style))
    story.append(Spacer(1, 0.3*inch))
    # Contact Information
    story.append(Paragraph("Contact Information", heading_style))
    contact = """
    <b>LF Systems</b><br/>
    A Division of RM Manifold Group Inc.<br/>
    <br/>
    <b>Website:</b> www.lfsystems.net<br/>
    <b>Phone:</b> (Contact your local representative)<br/>
    <b>Email:</b> info@lfsystems.net<br/>
    <br/>
    <i>Professional Dryer Exhaust Solutions Since 2008</i>
    """
    story.append(Paragraph(contact, body_style))
    # Footer with branding
    story.append(Spacer(1, 0.5*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER
    )
    footer = Paragraph(
        "This proposal is valid for 30 days from date of issue. "
        "Specifications subject to change based on final project requirements.<br/>"
        "© 2026 LF Systems by RM Manifold Group Inc. All rights reserved.",
        footer_style
    )
    story.append(footer)
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer
def generate_csi_specification(project_info, dryers, manifold_info, results):
    """Generate CSI specification as Word document"""
    doc = Document()
    # Set up document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("SECTION: 23 35 01")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("COMMERCIAL DRYER EXHAUST SYSTEM")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacing
    # PART 1: GENERAL
    heading1 = doc.add_heading("PART 1: GENERAL", level=1)
    # 1.01 SUMMARY
    doc.add_heading("1.01 SUMMARY", level=2)
    p = doc.add_paragraph()
    p.add_run("A. This system shall provide the required exhaust capacity for commercial clothes dryer applications. The following are components of the system:")
    dryers_list = doc.add_paragraph(style='List Number')
    dryers_list.add_run(f"LF Systems, Commercial Dryer Exhaust Fan Model DEF{results.get('selected_fan', {}).get('model', 'XX').replace('DEF', '')}, ETL-listed to UL STD 705.")
    controls_list = doc.add_paragraph(style='List Number')
    controls_list.add_run("LF Systems Model L150 Constant Pressure Controller")
    elec_list = doc.add_paragraph(style='List Number')
    elec_list.add_run("Electrical connections, by installing contractor.")
    # 1.02 RELATED REQUIREMENTS  
    doc.add_heading("1.02 RELATED REQUIREMENTS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Section 01 41 00, REGULATORY REQUIREMENTS")
    # 1.03 CODES AND STANDARDS
    doc.add_heading("1.03 CODES AND STANDARDS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. The following published specification standards apply to work in this section:")
    for code in ["UL -- Underwriters Laboratories", "ANSI Z223.1", "NFPA 211", "International Mechanical Code", "National Electrical Code"]:
        code_p = doc.add_paragraph(style='List Number')
        code_p.add_run(code)
    # 1.04 SUBMITTALS
    doc.add_heading("1.04 SUBMITTALS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Submittal documents shall be provided by the local manufacturer's representative.")
    p = doc.add_paragraph()
    p.add_run("B. The following shall be submitted to the Owner's representative:")
    submit_items = [
        "All relevant data sheets for the dryer exhaust fan, controller and related controls.",
        "ASHRAE duct design calculations.",
        "Wiring diagrams and installation manuals shall be submitted to the installing contractor prior to installation.",
        "Certification of listing for the actual application by an OSHA approved NRTL."
    ]
    for item in submit_items:
        item_p = doc.add_paragraph(style='List Number')
        item_p.add_run(item)
    # 1.05 QUALITY ASSURANCE
    doc.add_heading("1.05 QUALITY ASSURANCE", level=2)
    p = doc.add_paragraph()
    p.add_run(f"A. All listed dryer exhaust fans shall be assembled and marked at the facility indicated by the listing control number.")
    p = doc.add_paragraph()
    p.add_run("B. Exhaust system must be able to sense pressure in the venting system and maintain constant pressure.")
    p = doc.add_paragraph()
    p.add_run(f"C. System capacity scheduled shall be minimum {results.get('total_cfm', 0):.0f} CFM at {results.get('total_system_dp', 0):.2f} inches water column static pressure.")
    # 1.06 WARRANTY
    doc.add_heading("1.06 MANUFACTURER WARRANTY", level=2)
    p = doc.add_paragraph()
    p.add_run("A. All equipment is to be guaranteed against defects in materials and/or workmanship for a period of 24 months from the date of delivery to the construction site.")
    # PART 2: PRODUCTS
    doc.add_page_break()
    heading2 = doc.add_heading("PART 2: PRODUCTS", level=1)
    # 2.01 MANUFACTURER
    doc.add_heading("2.01 MANUFACTURER, DRYER EXHAUST SYSTEM", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Furnish LF Systems Commercial Dryer Exhaust System consisting of the following components:")
    fan_model = results.get('selected_fan', {}).get('model', 'DEFXX')
    components = [
        f"Model {fan_model} Inline dryer exhaust fan with design volume of {results.get('total_cfm', 0):.0f} CFM and design pressure of {results.get('total_system_dp', 0):.2f} inches water column as scheduled.",
        "Model L150 constant pressure controller with integrated display and alarm outputs.",
        "Current sensing relay (if required for application)."
    ]
    for comp in components:
        comp_p = doc.add_paragraph(style='List Number')
        comp_p.add_run(comp)
    # 2.02 EXHAUST FAN DESCRIPTION
    doc.add_heading("2.02 DESCRIPTION, INLINE DRYER EXHAUST FAN", level=2)
    fan_desc = [
        "The entire dryer exhaust fan shall be constructed of G90 galvanized steel with minimum .030 thickness. The fan shall be listed to UL STD 705 and shall bear the listed mark from an OSHA approved NRTL.",
        "The fan impeller shall be statically and dynamically balanced with permanently attached balancing weights of the same material as the impeller.",
        "The dryer exhaust fan shall be listed for 480°F exhaust gas temperatures and shall discharge gases as scheduled on drawings. The fan shall include integrated access panel.",
        "The fan motor shall be electronically commutated (EC), totally enclosed and outdoor rated with minimum efficiency of 75%. The motor shall provide variable speed control for pressure management."
    ]
    for idx, desc in enumerate(fan_desc, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # 2.03 CONTROLLER DESCRIPTION
    doc.add_heading("2.03 DESCRIPTION, L150 CONSTANT PRESSURE CONTROLLER", level=2)
    controller_desc = [
        "The L150 constant pressure controller shall monitor and control pressure in the dryer exhaust duct to maintain programmed setpoint pressure.",
        "The controller's pressure range shall be -1.00\" w.c. to +1.00\" w.c. with resolution of 0.01\" w.c.",
        "The controller shall provide visual LCD display, audible alarm, and dry contact alarm outputs for remote monitoring.",
        f"The controller shall maintain setpoint within +/-0.01\" w.c. Setpoint for this application: {results.get('manifold_dp', 0):.2f}\" w.c.",
        "The L150 shall include sensor verification, sleep mode, and adjustable PID control parameters."
    ]
    for idx, desc in enumerate(controller_desc, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # 2.04 PERFORMANCE
    doc.add_heading("2.04 PERFORMANCE REQUIREMENTS", level=2)
    perf = [
        "The dryer exhaust fan system shall reach setpoint within 15 seconds of initial demand.",
        f"Fan shall deliver minimum {results.get('total_cfm', 0):.0f} CFM at {results.get('total_system_dp', 0):.2f}\" w.c. system static pressure.",
        "The system shall include intelligent feedback signal to determine and control motor RPM.",
        "System shall operate quietly with vibration isolation as required."
    ]
    for idx, desc in enumerate(perf, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # 2.05 SEQUENCE OF OPERATION
    doc.add_heading("2.05 SEQUENCE OF OPERATION", level=2)
    seq = [
        "Upon demand signal, the L150 controller activates sensor verification. Once verified, the system controls fan speed to achieve setpoint pressure.",
        "As individual dryers call for heat, the L150 adjusts fan speed to maintain constant setpoint pressure in the exhaust manifold.",
        "When all dryers have satisfied, the controller engages sleep mode after programmable delay period.",
        "System includes integrated alarm functions for low pressure, high pressure, and sensor fault conditions."
    ]
    for idx, desc in enumerate(seq, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # 2.06 ELECTRICAL
    doc.add_heading("2.06 ELECTRICAL REQUIREMENTS", level=2)
    p = doc.add_paragraph()
    p.add_run("A. Power supply shall be:")
    elec = [
        "L150 pressure controller: 120VAC, single phase, 60Hz",
        "EC motor fan: 120VAC or 240VAC single phase, 60Hz (as scheduled)",
        "All wiring in accordance with National Electrical Code"
    ]
    for item in elec:
        elec_p = doc.add_paragraph(style='List Number')
        elec_p.add_run(item)
    # PART 3: EXECUTION
    doc.add_page_break()
    heading3 = doc.add_heading("PART 3: EXECUTION", level=1)
    # 3.01 INSTALLATION
    doc.add_heading("3.01 INSTALLATION", level=2)
    install = [
        "Complete structural, mechanical, and electrical connections in accordance with manufacturer's printed instructions.",
        f"Install {len(dryers)} dryer connectors to manifold duct of {manifold_info.get('diameter', 'XX')}\" diameter as shown on drawings.",
        f"Install manifold duct system of {manifold_info.get('length', 'XX')} feet total length with fittings as scheduled.",
        f"Mount {fan_model} exhaust fan in location shown on drawings with proper vibration isolation and support.",
        "Install L150 controller in accessible location for programming and service. Mount pressure sensor in manifold as detailed.",
        "Low voltage control wiring shall be minimum 18 AWG shielded wire. Maintain separation from line voltage wiring.",
        "Provide electrical disconnects and overcurrent protection as required by NEC."
    ]
    for idx, desc in enumerate(install, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # 3.02 STARTUP
    doc.add_heading("3.02 STARTUP AND COMMISSIONING", level=2)
    startup = [
        "System startup shall be performed by qualified LF Systems technician or authorized representative.",
        "Verify all electrical connections, duct connections, and sensor locations prior to energizing system.",
        f"Program L150 controller with setpoint pressure of {results.get('manifold_dp', 0):.2f}\" w.c. and verify sensor calibration.",
        "Test system operation with all connected dryers and verify pressure control performance.",
        "Provide training to owner's maintenance personnel on system operation, troubleshooting, and routine maintenance."
    ]
    for idx, desc in enumerate(startup, 1):
        p = doc.add_paragraph()
        p.add_run(f"{chr(64+idx)}. {desc}")
    # Footer
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("LF SYSTEMS by RM Manifold Group Inc.\n")
    footer_run.bold = True
    footer_run.font.size = Pt(12)
    footer_run2 = footer_p.add_run("www.lfsystems.net\n")
    footer_run2.font.size = Pt(10)
    footer_run3 = footer_p.add_run(f"Project: {project_info.get('name', 'Commercial Dryer Exhaust')}\n")
    footer_run3.font.size = Pt(10)
    footer_run4 = footer_p.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    footer_run4.font.size = Pt(10)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
def main():
    st.set_page_config(
        page_title="DREX - Dryer Exhaust Calculator | LF Systems", 
        layout="wide", 
        page_icon="🧺"  # Laundry/dryer icon
    )
    initialize_session_state()
    
    # Custom CSS for LF Systems branding
    st.markdown(f"""
    <style>
        .stMetric {{
            background-color: #f8f9fa;
            padding: 10px;
            border-left: 4px solid {BRAND_PRIMARY};
            border-radius: 4px;
        }}
        .stButton>button {{
            background-color: {BRAND_SECONDARY};
            color: white;
            font-weight: 500;
        }}
        .stButton>button:hover {{
            background-color: {BRAND_PRIMARY};
        }}
        h1 {{
            color: {BRAND_PRIMARY};
        }}
    </style>
    """, unsafe_allow_html=True)
    
    # Header with LF Systems logo
    try:
        col1, col2 = st.columns([1, 3])
        with col1:
            st.image("lf_systems_logo.jpg", width=300)
        with col2:
            st.title("DREX - Commercial Dryer Exhaust Calculator")
            st.markdown("*Professional Dryer Exhaust System Design & Analysis*")
            st.markdown("🌐 [www.lfsystems.net](http://www.lfsystems.net)")
    except:
        # Fallback if logo not found
        col1, col2 = st.columns([1, 5])
        with col1:
            st.markdown("# 🧺")
        with col2:
            st.title("DREX - Commercial Dryer Exhaust Calculator")
            st.markdown("**By LF Systems** - *A Division of RM Manifold Group Inc.*")
            st.markdown("🌐 [www.lfsystems.net](http://www.lfsystems.net)")
    st.markdown("---")
    
    # Sidebar for navigation
    with st.sidebar:
        st.header("Navigation")
        if st.button("🏠 Start Over"):
            st.session_state.clear()
            st.rerun()
        
        st.markdown("---")
        st.markdown("### Current Progress")
        steps = {
            'welcome': '✓ Welcome',
            'project_info': '✓ Project Info' if st.session_state.step != 'welcome' else '○ Project Info',
            'dryer_input': '✓ Dryer Input' if st.session_state.step not in ['welcome', 'project_info'] else '○ Dryer Input',
            'manifold_input': '✓ Manifold Input' if st.session_state.step not in ['welcome', 'project_info', 'dryer_input'] else '○ Manifold Input',
            'results': '✓ Results' if st.session_state.step == 'results' else '○ Results'
        }
        for step_name, label in steps.items():
            st.markdown(f"**{label}**")
    
    # Main content based on current step
    if st.session_state.step == 'welcome':
        show_welcome_screen()
    elif st.session_state.step == 'project_info':
        show_project_info_screen()
    elif st.session_state.step == 'dryer_input':
        show_dryer_input_screen()
    elif st.session_state.step == 'manifold_input':
        show_manifold_input_screen()
    elif st.session_state.step == 'results':
        show_results_screen()

def show_welcome_screen():
    """Display welcome screen"""
    st.header("Welcome to DREX!")
    st.markdown("""
    ### Commercial Dryer Exhaust System Sizing Calculator
    
    This calculator will help you design a compliant commercial dryer exhaust system by:
    
    - 📊 Calculating pressure losses through dryer connectors and manifold ductwork
    - ⚠️ Identifying potential issues with excessive pressure loss
    - 🔧 Optimizing manifold diameter for maximum efficiency
    - 📈 Selecting the appropriate DEF exhaust fan
    - 📄 Generating professional PDF reports and CSI specifications
    
    #### System Capabilities:
    - Up to 20 individual commercial dryers
    - Up to 4 different dryer capacities
    - Automatic or manual manifold sizing
    - Lint collector accommodation
    - Complete documentation package
    
    #### Design Criteria:
    - Maximum dryer connector loss: 0.25 IN WC (warning issued if exceeded)
    - Maximum manifold loss: 1.0 IN WC
    - Lateral tees only (no 90° tees permitted)
    
    Click **Next** to begin entering your project information.
    """)
    
    if st.button("Next ➡️", type="primary"):
        st.session_state.step = 'project_info'
        st.rerun()

def show_project_info_screen():
    """Display project information input screen"""
    st.header("Project Information")
    
    with st.form("project_info_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            project_name = st.text_input("Project Name*", value=st.session_state.project_info.get('name', ''))
            zip_code = st.text_input("Zip Code*", value=st.session_state.project_info.get('zip', ''), max_chars=5)
            user_name = st.text_input("Your Name*", value=st.session_state.project_info.get('user_name', ''))
        
        with col2:
            user_email = st.text_input("Your Email*", value=st.session_state.project_info.get('user_email', ''))
            
            # Show location info if zip code is available
            if zip_code and len(zip_code) == 5:
                location_info = get_elevation_from_zip(zip_code)
                if location_info['success']:
                    st.info(f"📍 {location_info['city']}, {location_info['state']}")
                    st.info(f"⛰️ Elevation: {location_info['elevation']:.0f} feet")
        
        submitted = st.form_submit_button("Save and Continue ➡️", type="primary")
        
        if submitted:
            if not all([project_name, zip_code, user_name, user_email]):
                st.error("Please fill in all required fields.")
            elif len(zip_code) != 5 or not zip_code.isdigit():
                st.error("Please enter a valid 5-digit zip code.")
            else:
                location_info = get_elevation_from_zip(zip_code)
                
                st.session_state.project_info = {
                    'name': project_name,
                    'zip': zip_code,
                    'user_name': user_name,
                    'user_email': user_email,
                    'city': location_info.get('city', 'Unknown') if location_info['success'] else 'Unknown',
                    'state': location_info.get('state', 'Unknown') if location_info['success'] else 'Unknown',
                    'elevation': location_info.get('elevation', 0) if location_info['success'] else 0
                }
                st.session_state.step = 'dryer_input'
                st.rerun()

def show_dryer_input_screen():
    """Display dryer input screen"""
    st.header("Dryer Information")
    st.markdown(f"**Project:** {st.session_state.project_info.get('name', 'N/A')}")
    
    # Show existing dryers
    if st.session_state.dryers:
        st.subheader(f"Current Dryers ({len(st.session_state.dryers)}/20)")
        
        for idx, dryer in enumerate(st.session_state.dryers):
            with st.expander(f"Dryer #{idx + 1} - {dryer['cfm']} CFM"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.write(f"**CFM:** {dryer['cfm']}")
                    st.write(f"**Outlet Diameter:** {dryer['outlet_diameter']} inches")
                with col2:
                    st.write(f"**Connector Length:** {dryer['connector_length']} feet")
                    st.write(f"**Total K-value:** {dryer['total_k']:.2f}")
                with col3:
                    if st.button(f"Remove Dryer #{idx + 1}", key=f"remove_{idx}"):
                        st.session_state.dryers.pop(idx)
                        st.rerun()
        
        st.markdown("---")
    
    # Add new dryer form
    if len(st.session_state.dryers) < 20:
        st.subheader("Add New Dryer")
        
        with st.form("add_dryer_form"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                cfm = st.number_input("Dryer CFM*", min_value=50, max_value=5000, value=200, step=50)
                outlet_diameter = st.number_input("Dryer Outlet Diameter (inches)*", min_value=3, max_value=12, value=4, step=1)
            
            with col2:
                connector_length = st.number_input("Connector Length (feet)*", min_value=1, max_value=100, value=10, step=1)
                additional_k = st.number_input("Additional K-value (lint collector, etc.)", min_value=0.0, max_value=10.0, value=0.0, step=0.1)
            
            with col3:
                st.markdown("**Connector Fittings**")
                num_90_elbows = st.number_input("90° Elbows", min_value=0, max_value=10, value=2, step=1)
                num_45_elbows = st.number_input("45° Elbows", min_value=0, max_value=10, value=0, step=1)
                num_30_elbows = st.number_input("30° Elbows", min_value=0, max_value=10, value=0, step=1)
                num_lateral_tees = st.number_input("Lateral Tees", min_value=0, max_value=5, value=1, step=1)
            
            col_a, col_b, col_c = st.columns(3)
            
            with col_a:
                copy_count = st.number_input("Number of copies to add", min_value=1, max_value=20, value=1, step=1)
            
            submitted = st.form_submit_button("Add Dryer(s)", type="primary")
            
            if submitted:
                # Calculate total K-value
                k_total = (
                    num_90_elbows * K_VALUES["90° Elbow"] +
                    num_45_elbows * K_VALUES["45° Elbow"] +
                    num_30_elbows * K_VALUES["30° Elbow"] +
                    num_lateral_tees * K_VALUES["Lateral Tee"] +
                    K_VALUES["Entry (flush)"] +
                    K_VALUES["Exit"] +
                    additional_k
                )
                
                # Create fittings summary
                fittings = []
                if num_90_elbows > 0:
                    fittings.append(f"{num_90_elbows}x 90° Elbow")
                if num_45_elbows > 0:
                    fittings.append(f"{num_45_elbows}x 45° Elbow")
                if num_30_elbows > 0:
                    fittings.append(f"{num_30_elbows}x 30° Elbow")
                if num_lateral_tees > 0:
                    fittings.append(f"{num_lateral_tees}x Lateral Tee")
                fittings.append("Entry + Exit")
                fittings_summary = ", ".join(fittings)
                
                # Calculate velocity and pressure loss
                velocity = calculate_velocity(cfm, outlet_diameter)
                dp = calculate_duct_pressure_loss(connector_length, outlet_diameter, velocity, k_total)
                
                # Create dryer object
                dryer_obj = {
                    'cfm': cfm,
                    'outlet_diameter': outlet_diameter,
                    'connector_length': connector_length,
                    'additional_k': additional_k,
                    'total_k': k_total,
                    'connector_fittings_summary': fittings_summary,
                    'connector_velocity': velocity,
                    'connector_dp': dp
                }
                
                # Add copies
                for _ in range(min(copy_count, 20 - len(st.session_state.dryers))):
                    st.session_state.dryers.append(dryer_obj.copy())
                
                st.rerun()
    
    else:
        st.warning("Maximum of 20 dryers reached.")
    
    # Navigation buttons
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.button("⬅️ Back to Project Info"):
            st.session_state.step = 'project_info'
            st.rerun()
    
    with col3:
        if st.session_state.dryers:
            if st.button("Continue to Manifold ➡️", type="primary"):
                st.session_state.step = 'manifold_input'
                st.rerun()
        else:
            st.info("Add at least one dryer to continue")

def show_manifold_input_screen():
    """Display manifold input screen"""
    st.header("Manifold Duct Information")
    st.markdown(f"**Project:** {st.session_state.project_info.get('name', 'N/A')}")
    st.markdown(f"**Total Dryers:** {len(st.session_state.dryers)}")
    
    total_cfm = sum(d['cfm'] for d in st.session_state.dryers)
    st.markdown(f"**Total System CFM:** {total_cfm} CFM")
    
    # Optimization choice outside form for dynamic update
    st.subheader("Manifold Configuration")
    optimize_diameter = st.checkbox("✅ Optimize manifold diameter automatically (target 1.0 IN WC max)", value=True, key="opt_check")
    
    if not optimize_diameter:
        st.info("💡 Manual diameter mode: You will specify the diameter below")
    else:
        st.info("🔧 Optimization mode: Diameter will be calculated automatically")
    
    st.markdown("---")
    
    with st.form("manifold_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            manifold_length = st.number_input("Manifold Length (feet)*", min_value=1, max_value=500, value=50, step=5)
            
            # Show diameter input based on checkbox state
            if not optimize_diameter:
                manifold_diameter = st.number_input(
                    "Manifold Diameter (inches)*", 
                    min_value=6, 
                    max_value=36, 
                    value=12, 
                    step=1,
                    help="Standard sizes: 6, 8, 10, 12, 14, 16, 18, 20, 22, 24, 26, 28, 30, 32, 34, 36 inches"
                )
            else:
                manifold_diameter = None
        
        with col2:
            st.markdown("**Manifold Fittings**")
            st.info("Note: Do NOT count straight-through tees. Only count lateral tees (change of direction).")
            num_90_elbows = st.number_input("90° Elbows", min_value=0, max_value=20, value=2, step=1, key="m_90")
            num_45_elbows = st.number_input("45° Elbows", min_value=0, max_value=20, value=0, step=1, key="m_45")
            num_30_elbows = st.number_input("30° Elbows", min_value=0, max_value=20, value=0, step=1, key="m_30")
            num_lateral_tees = st.number_input("Lateral Tees (change of direction only)", min_value=0, max_value=20, value=0, step=1, key="m_tee")
            additional_k_manifold = st.number_input("Additional K-value (lint collector, etc.)", min_value=0.0, max_value=10.0, value=0.0, step=0.1, key="m_addk")
        
        submitted = st.form_submit_button("Calculate System ➡️", type="primary")
        
        if submitted:
            # Calculate total K-value for manifold (no straight-through tees)
            k_total_manifold = (
                num_90_elbows * K_VALUES["90° Elbow"] +
                num_45_elbows * K_VALUES["45° Elbow"] +
                num_30_elbows * K_VALUES["30° Elbow"] +
                num_lateral_tees * K_VALUES["Lateral Tee"] +
                K_VALUES["Entry (flush)"] +
                K_VALUES["Exit"] +
                additional_k_manifold
            )
            
            # Create fittings summary
            fittings = []
            if num_90_elbows > 0:
                fittings.append(f"{num_90_elbows}x 90° Elbow")
            if num_45_elbows > 0:
                fittings.append(f"{num_45_elbows}x 45° Elbow")
            if num_30_elbows > 0:
                fittings.append(f"{num_30_elbows}x 30° Elbow")
            if num_lateral_tees > 0:
                fittings.append(f"{num_lateral_tees}x Lateral Tee")
            fittings.append("Entry + Exit")
            fittings_summary = ", ".join(fittings)
            
            # Optimize or use specified diameter
            if optimize_diameter:
                opt_diameter, velocity, dp = optimize_manifold_diameter(total_cfm, manifold_length, k_total_manifold)
                manifold_diameter = opt_diameter
            else:
                velocity = calculate_velocity(total_cfm, manifold_diameter)
                dp = calculate_duct_pressure_loss(manifold_length, manifold_diameter, velocity, k_total_manifold)
            
            # Save manifold info
            st.session_state.manifold_info = {
                'length': manifold_length,
                'diameter': manifold_diameter,
                'optimize': optimize_diameter,
                'fittings_summary': fittings_summary,
                'additional_k': additional_k_manifold,
                'total_k': k_total_manifold
            }
            
            # Calculate results
            worst_connector_dp = max(d['connector_dp'] for d in st.session_state.dryers)
            total_system_dp = worst_connector_dp + dp
            
            # Select fan
            suitable_fans = select_def_fan(total_cfm, total_system_dp)
            
            st.session_state.calculation_results = {
                'total_cfm': total_cfm,
                'worst_connector_dp': worst_connector_dp,
                'manifold_dp': dp,
                'manifold_velocity': velocity,
                'manifold_total_k': k_total_manifold,
                'total_system_dp': total_system_dp,
                'suitable_fans': suitable_fans,
                'selected_fan': suitable_fans[0] if suitable_fans else None
            }
            
            st.session_state.step = 'results'
            st.rerun()
    
    # Navigation
    st.markdown("---")
    if st.button("⬅️ Back to Dryer Input"):
        st.session_state.step = 'dryer_input'
        st.rerun()

def show_results_screen():
    """Display results and generate reports"""
    st.header("System Calculation Results")
    st.markdown(f"**Project:** {st.session_state.project_info.get('name', 'N/A')}")
    
    results = st.session_state.calculation_results
    
    # System Summary
    st.subheader("📊 System Summary")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total System CFM", f"{results['total_cfm']:.0f}")
    with col2:
        st.metric("Manifold Diameter", f"{st.session_state.manifold_info['diameter']}\"")
    with col3:
        st.metric("Manifold Velocity", f"{results['manifold_velocity']:.0f} FPM")
    with col4:
        st.metric("Total System ΔP", f"{results['total_system_dp']:.3f} IN WC")
    
    st.markdown("---")
    
    # Pressure Loss Breakdown
    st.subheader("📉 Pressure Loss Breakdown")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.metric(
            "Worst Case Connector Loss",
            f"{results['worst_connector_dp']:.3f} IN WC",
            delta=f"{results['worst_connector_dp'] - DRYER_CONNECTOR_MAX_DP:.3f} IN WC" if results['worst_connector_dp'] > DRYER_CONNECTOR_MAX_DP else None,
            delta_color="inverse"
        )
    
    with col2:
        st.metric(
            "Manifold Loss",
            f"{results['manifold_dp']:.3f} IN WC",
            delta=f"{results['manifold_dp'] - MANIFOLD_MAX_DP:.3f} IN WC" if results['manifold_dp'] > MANIFOLD_MAX_DP else None,
            delta_color="inverse"
        )
    
    # Warnings
    if results['worst_connector_dp'] > DRYER_CONNECTOR_MAX_DP:
        st.warning(f"⚠️ **WARNING:** Worst case dryer connector pressure loss ({results['worst_connector_dp']:.3f} IN WC) exceeds recommended maximum of {DRYER_CONNECTOR_MAX_DP} IN WC. Consider reducing connector length, increasing diameter, or adding a booster fan.")
    
    if results['manifold_dp'] > MANIFOLD_MAX_DP:
        st.error(f"❌ **ERROR:** Manifold pressure loss ({results['manifold_dp']:.3f} IN WC) exceeds maximum of {MANIFOLD_MAX_DP} IN WC. Increase manifold diameter.")
    
    st.markdown("---")
    
    # Fan Selection
    st.subheader("🔧 Fan Selection")
    
    if results['selected_fan']:
        fan = results['selected_fan']
        
        # Fan information and curve plot side by side
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.success(f"**Recommended Fan:** {fan['model']}")
            st.write(f"**Available CFM at Operating Point:** {fan['available_cfm']:.0f} CFM")
            st.write(f"**Design Margin:** {fan['margin']:.1f}%")
            st.write(f"**Maximum Fan Capacity:** {fan['max_cfm']:.0f} CFM @ {fan['max_sp']:.2f} IN WC")
            
            # Show alternative fans if available
            if len(results['suitable_fans']) > 1:
                with st.expander("View Alternative Fan Options"):
                    for alt_fan in results['suitable_fans'][1:4]:  # Show up to 3 alternatives
                        st.write(f"**{alt_fan['model']}** - {alt_fan['available_cfm']:.0f} CFM @ operating point ({alt_fan['margin']:.1f}% margin)")
        
        with col2:
            # Show fan curve plot
            st.markdown("**Fan Performance Curve:**")
            fig = plot_fan_and_system_curves(
                fan['model'],
                results['total_cfm'],
                results['total_system_dp'],
                results['manifold_dp']  # Manifold pressure only for system curve
            )
            st.pyplot(fig)
            plt.close()
    
    else:
        st.error("❌ No suitable DEF fan found for the system requirements. System pressure loss may be too high or CFM too low.")
    
    st.markdown("---")
    
    # Generate Reports
    st.subheader("📄 Generate Reports")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 PDF Technical Report", type="primary", use_container_width=True):
            pdf_buffer = generate_pdf_report(
                st.session_state.project_info,
                st.session_state.dryers,
                st.session_state.manifold_info,
                results
            )
            
            st.download_button(
                label="📥 Download PDF Report",
                data=pdf_buffer,
                file_name=f"LF_Systems_DREX_Report_{st.session_state.project_info.get('name', 'Project').replace(' ', '_')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
    
    with col2:
        if st.button("📋 CSI Specification", use_container_width=True):
            csi_spec = generate_csi_specification(
                st.session_state.project_info,
                st.session_state.dryers,
                st.session_state.manifold_info,
                results
            )
            
            st.download_button(
                label="📥 Download CSI Spec",
                data=csi_spec,
                file_name=f"LF_Systems_CSI_Spec_{st.session_state.project_info.get('name', 'Project').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    # Navigation
    st.markdown("---")
    col1, col2 = st.columns([1, 1])
    
    with col1:
        if st.button("⬅️ Back to Manifold Input"):
            st.session_state.step = 'manifold_input'
            st.rerun()
    
    with col2:
        if st.button("🏠 Start New Project"):
            st.session_state.clear()
            st.rerun()

if __name__ == "__main__":
    main()
